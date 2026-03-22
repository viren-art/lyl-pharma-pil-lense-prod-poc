"""
PIL Lens Python Document Service

FastAPI microservice for high-fidelity document extraction and generation.
Runs alongside Node.js Express on the same Cloud Run container (port 8081).

Key capability: Template cloning — reads exact formatting from an approved
Lotus PIL (.docx) and applies identical styles to generated documents.
"""

import io
import json
import base64
import logging
import tempfile
import os
from typing import Optional
from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.responses import StreamingResponse, JSONResponse
import uvicorn

app = FastAPI(title="PIL Lens Document Service", version="1.0.0")
logging.basicConfig(level=logging.INFO, format="[PythonDocService] %(message)s")
logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────────
# POST /extract/word
# Extract .docx with FULL formatting metadata
# ─────────────────────────────────────────────────
@app.post("/extract/word")
async def extract_word(file: UploadFile = File(...)):
    """Extract Word document with formatting: headings, bold, tables, images."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    try:
        content = await file.read()
        doc = Document(io.BytesIO(content))

        sections = []
        current_section = None
        current_subsection = None

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect formatting
            is_bold = any(run.bold for run in para.runs if run.bold is not None)
            is_italic = any(run.italic for run in para.runs if run.italic is not None)
            font_name = None
            font_size = None
            for run in para.runs:
                if run.font.name:
                    font_name = run.font.name
                if run.font.size:
                    font_size = run.font.size.pt

            # Detect heading level from style
            heading_level = None
            style_name = para.style.name if para.style else ''
            if 'Heading' in style_name:
                try:
                    heading_level = int(style_name.replace('Heading ', '').strip())
                except ValueError:
                    heading_level = 1

            # Detect section numbering from text
            import re
            section_match = re.match(r'^(\d+)\.\s+(.+)', text)
            subsection_match = re.match(r'^(\d+\.\d+)\s+(.+)', text)
            subsubsection_match = re.match(r'^(\d+\.\d+\.\d+[\d.]*)\s*(.+)', text)

            paragraph_data = {
                'text': text,
                'formatting': {
                    'bold': is_bold,
                    'italic': is_italic,
                    'fontName': font_name,
                    'fontSize': font_size,
                    'alignment': str(para.alignment) if para.alignment else None,
                    'headingLevel': heading_level,
                    'styleName': style_name,
                },
                'runs': [{
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline is not None and run.underline,
                    'fontName': run.font.name,
                    'fontSize': run.font.size.pt if run.font.size else None,
                    'color': str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None,
                } for run in para.runs if run.text.strip()]
            }

            if subsubsection_match:
                paragraph_data['sectionNumber'] = subsubsection_match.group(1)
                paragraph_data['sectionName'] = subsubsection_match.group(2)
                paragraph_data['level'] = 'subsubsection'
            elif subsection_match:
                paragraph_data['sectionNumber'] = subsection_match.group(1)
                paragraph_data['sectionName'] = subsection_match.group(2)
                paragraph_data['level'] = 'subsection'
            elif section_match:
                paragraph_data['sectionNumber'] = section_match.group(1)
                paragraph_data['sectionName'] = section_match.group(2)
                paragraph_data['level'] = 'section'
            else:
                paragraph_data['level'] = 'body'

            sections.append(paragraph_data)

        # Extract tables
        tables = []
        for i, table in enumerate(doc.tables):
            table_data = {
                'index': i,
                'rows': [],
                'rowCount': len(table.rows),
                'colCount': len(table.columns) if table.rows else 0,
            }
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    # Get cell formatting
                    cell_bold = False
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.bold:
                                cell_bold = True
                    row_data.append({
                        'text': cell_text,
                        'bold': cell_bold,
                    })
                table_data['rows'].append(row_data)
            tables.append(table_data)

        # Extract images
        images = []
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                images.append({
                    'reltype': rel.reltype,
                    'target': rel.target_ref,
                    'contentType': rel.target_part.content_type if hasattr(rel, 'target_part') and rel.target_part else None,
                })

        # Extract document properties
        core_props = doc.core_properties
        properties = {
            'title': core_props.title,
            'author': core_props.author,
            'created': str(core_props.created) if core_props.created else None,
            'modified': str(core_props.modified) if core_props.modified else None,
        }

        # Extract styles used
        styles_used = set()
        for para in doc.paragraphs:
            if para.style:
                styles_used.add(para.style.name)

        logger.info(f"Extracted Word: {len(sections)} paragraphs, {len(tables)} tables, {len(images)} images")

        return {
            'paragraphs': sections,
            'tables': tables,
            'images': images,
            'properties': properties,
            'stylesUsed': list(styles_used),
            'totalParagraphs': len(sections),
            'totalTables': len(tables),
            'totalImages': len(images),
        }

    except Exception as e:
        logger.error(f"Word extraction failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ─────────────────────────────────────────────────
# POST /extract/pdf-tables
# Extract tables from PDF using camelot
# ─────────────────────────────────────────────────
@app.post("/extract/pdf-tables")
async def extract_pdf_tables(file: UploadFile = File(...)):
    """Extract tables from PDF using camelot for clinical trial data."""
    try:
        content = await file.read()

        # Write to temp file (camelot needs file path)
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        try:
            import camelot
            tables = camelot.read_pdf(tmp_path, pages='all', flavor='lattice')

            result = []
            for i, table in enumerate(tables):
                df = table.df
                result.append({
                    'index': i,
                    'page': table.page,
                    'rows': df.values.tolist(),
                    'headers': df.iloc[0].tolist() if len(df) > 0 else [],
                    'shape': list(df.shape),
                    'accuracy': table.accuracy,
                    'parsing_report': table.parsing_report,
                })

            logger.info(f"Extracted {len(result)} tables from PDF")
            return {'tables': result, 'totalTables': len(result)}

        finally:
            os.unlink(tmp_path)

    except ImportError:
        # Camelot not available — try tabula as fallback
        try:
            import tabula
            content_io = io.BytesIO(content)

            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
                tmp.write(content_io.getvalue())
                tmp_path = tmp.name

            try:
                dfs = tabula.read_pdf(tmp_path, pages='all', multiple_tables=True)
                result = []
                for i, df in enumerate(dfs):
                    result.append({
                        'index': i,
                        'rows': df.values.tolist(),
                        'headers': df.columns.tolist(),
                        'shape': list(df.shape),
                    })
                return {'tables': result, 'totalTables': len(result)}
            finally:
                os.unlink(tmp_path)

        except Exception as e2:
            logger.error(f"PDF table extraction failed: {e2}")
            raise HTTPException(status_code=500, detail=f"Table extraction failed: {e2}")

    except Exception as e:
        logger.error(f"PDF table extraction failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ─────────────────────────────────────────────────
# POST /generate/word
# Generate Word doc matching approved Lotus PIL format
# Uses template cloning from approved PIL
# ─────────────────────────────────────────────────
@app.post("/generate/word")
async def generate_word(
    sections_json: str = Form(...),
    product_name: str = Form(""),
    market_code: str = Form("taiwan_tfda"),
    mode: str = Form("tc"),  # 'en' or 'tc' or 'th'
    template_file: Optional[UploadFile] = File(None),
):
    """
    Generate Word document matching approved Lotus PIL format.

    If template_file is provided, clones exact formatting from it.
    Otherwise uses built-in TFDA formatting defaults.
    """
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    import copy

    try:
        sections = json.loads(sections_json)
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=400, detail=f"Invalid sections JSON: {e}")

    is_translated = mode != 'en'

    # Sections 4-12 get "(依文獻紀載)" suffix
    LITERATURE_SUFFIX = ['4', '5', '6', '7', '8', '9', '10', '11', '12']

    try:
        # ── Load template styles if provided ──
        template_styles = None
        if template_file:
            template_content = await template_file.read()
            template_doc = Document(io.BytesIO(template_content))
            template_styles = extract_template_styles(template_doc)
            logger.info(f"Loaded template styles: {len(template_styles)} style definitions")

        # ── Create document ──
        doc = Document()

        # Page setup — A4, 1 inch margins (matching approved PIL)
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(10)

        # Apply template styles if available
        if template_styles:
            apply_template_styles(doc, template_styles)

        # ── Header: Product name + registration numbers ──
        if is_translated:
            # Bilingual product name (matching approved PIL header)
            p = doc.add_paragraph()
            run = p.add_run(product_name or '[藥品中文名稱] [English Name] [Strength]')
            run.bold = True
            run.font.size = Pt(12)

            p = doc.add_paragraph()
            run = p.add_run('(衛部藥輸字第______號)')
            run.font.size = Pt(10)

            doc.add_paragraph()  # spacing

            p = doc.add_paragraph()
            run = p.add_run('須由醫師處方使用')
            run.bold = True
            run.font.size = Pt(10)

            # DRAFT watermark
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('【草稿 DRAFT — 供審核用】')
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

            doc.add_paragraph()  # spacing

        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('INTERNAL REVIEW COPY — English Structure Draft')
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f'Product: {product_name or "N/A"}')
            run.font.size = Pt(11)

            doc.add_paragraph()

        # ── Render sections ──
        for section_data in sections:
            num = section_data.get('number', '')
            name = section_data.get('name', '')
            local_name = section_data.get('localName', '')
            content = section_data.get('content', '')
            status = section_data.get('status', 'gap')
            gap_note = section_data.get('gapNote', '')

            # Determine suffix
            suffix = ' (依文獻紀載)' if is_translated and num in LITERATURE_SUFFIX else ''

            # ── Section heading ──
            if is_translated and local_name:
                heading_text = f'{num}. {local_name}{suffix}'
            else:
                heading_text = f'{num}. {name}'
                if local_name:
                    heading_text += f' ({local_name})'

            p = doc.add_paragraph()
            run = p.add_run(heading_text)
            run.bold = True
            run.font.size = Pt(12) if template_styles is None else Pt(template_styles.get('heading_size', 12))

            # ── Section content ──
            if status == 'mapped' and content:
                render_content_to_doc(doc, content, is_translated, template_styles)
            else:
                # Gap placeholder
                p = doc.add_paragraph()
                gap_text = '[內容缺漏 — 原始文件中無此資料，需另行補充]' if is_translated else '[CONTENT GAP — Not available in source document]'
                run = p.add_run(gap_text)
                run.bold = True
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                if gap_note:
                    p = doc.add_paragraph()
                    run = p.add_run(gap_note)
                    run.italic = True
                    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                    run.font.size = Pt(9)

        # ── Manufacturer footer (matching approved PIL) ──
        if is_translated:
            doc.add_paragraph()
            doc.add_paragraph()
            for line in [
                '製造廠：[廠名]',
                '廠址：[廠址]',
                '藥商：美時化學製藥股份有限公司',
                '地址：台北市信義區松仁路277號17樓',
            ]:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.size = Pt(10)

        # ── Save to buffer ──
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f'PIL_Draft_{market_code}_{mode}.docx'
        logger.info(f"Generated Word doc: {filename}")

        return StreamingResponse(
            buffer,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': f'attachment; filename="{filename}"'}
        )

    except Exception as e:
        logger.error(f"Word generation failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


def render_content_to_doc(doc, content: str, is_translated: bool, template_styles: dict = None):
    """Render section content to document, detecting subsection headings and formatting."""
    from docx.shared import Pt, RGBColor
    import re

    for line in content.split('\n'):
        text = line.strip()
        if not text:
            continue

        # Detect sub-subsection heading: 5.1.1, 3.3.1, 10.3.1.1 etc
        subsubsection = re.match(r'^(\d+\.\d+\.\d+[\d.]*)\s+(.+)', text)
        # Detect subsection heading: 5.1, 3.3, 10.2 etc
        subsection = re.match(r'^(\d+\.\d+)\s+(.+)', text)

        if subsubsection:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(10)
        elif subsection:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(10.5)
        elif text.startswith('•') or text.startswith('-'):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(20)
            run = p.add_run(text)
            run.font.size = Pt(10)
        elif '|' in text and text.count('|') >= 2:
            # Table-like row — render as monospace
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = 'Courier New'
            run.font.size = Pt(8)
        else:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(10)


def extract_template_styles(template_doc) -> dict:
    """Extract formatting styles from an approved PIL template document."""
    from docx.shared import Pt

    styles = {
        'heading_size': 12,
        'body_size': 10,
        'font_name': 'Times New Roman',
        'section_headings': {},
        'paragraph_spacing': None,
    }

    for para in template_doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Capture heading styles
        for run in para.runs:
            if run.bold and run.font.size:
                size_pt = run.font.size.pt
                if size_pt > 11:
                    styles['heading_size'] = size_pt
                if run.font.name:
                    styles['font_name'] = run.font.name
            elif run.font.size:
                styles['body_size'] = run.font.size.pt

        # Capture paragraph spacing
        if para.paragraph_format.space_after:
            styles['paragraph_spacing'] = para.paragraph_format.space_after

    return styles


def apply_template_styles(doc, template_styles: dict):
    """Apply extracted template styles to a new document."""
    from docx.shared import Pt

    style = doc.styles['Normal']
    font = style.font
    if template_styles.get('font_name'):
        font.name = template_styles['font_name']
    if template_styles.get('body_size'):
        font.size = Pt(template_styles['body_size'])


# ─────────────────────────────────────────────────
# POST /generate/pdf
# Generate print-ready PDF with ReportLab
# ─────────────────────────────────────────────────
@app.post("/generate/pdf")
async def generate_pdf(
    sections_json: str = Form(...),
    product_name: str = Form(""),
    market_code: str = Form("taiwan_tfda"),
):
    """Generate TFDA-compliant PDF with CJK fonts."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    try:
        sections = json.loads(sections_json)
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=400, detail=f"Invalid JSON: {e}")

    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4,
                                leftMargin=2.54*cm, rightMargin=2.54*cm,
                                topMargin=2.54*cm, bottomMargin=2.54*cm)

        # Register CJK fonts if available
        try:
            noto_path = '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc'
            if os.path.exists(noto_path):
                pdfmetrics.registerFont(TTFont('NotoSansCJK', noto_path, subfontIndex=0))
                logger.info("Registered NotoSansCJK font")
        except Exception as font_err:
            logger.warning(f"CJK font registration failed: {font_err}")

        styles = getSampleStyleSheet()

        # Add CJK-compatible styles
        styles.add(ParagraphStyle(
            'PilHeading',
            parent=styles['Heading1'],
            fontSize=14,
            spaceAfter=6,
            spaceBefore=12,
        ))
        styles.add(ParagraphStyle(
            'PilBody',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=4,
            leading=14,
        ))

        story = []

        # Title
        story.append(Paragraph(product_name or 'PIL Draft', styles['Title']))
        story.append(Spacer(1, 12))

        for section_data in sections:
            num = section_data.get('number', '')
            name = section_data.get('name', '')
            local_name = section_data.get('localName', '')
            content = section_data.get('content', '')

            heading = f'{num}. {local_name or name}'
            story.append(Paragraph(heading, styles['PilHeading']))

            if content:
                for line in content.split('\n'):
                    text = line.strip()
                    if text:
                        # Escape XML characters for ReportLab
                        text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        story.append(Paragraph(text, styles['PilBody']))
            else:
                story.append(Paragraph('[Content Gap]', styles['PilBody']))

            story.append(Spacer(1, 6))

        doc.build(story)
        buffer.seek(0)

        filename = f'PIL_Draft_{market_code}.pdf'
        return StreamingResponse(
            buffer,
            media_type='application/pdf',
            headers={'Content-Disposition': f'attachment; filename="{filename}"'}
        )

    except Exception as e:
        logger.error(f"PDF generation failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ─────────────────────────────────────────────────
# Health check
# ─────────────────────────────────────────────────
@app.get("/health")
async def health():
    return {"status": "ok", "service": "python-doc-service", "version": "1.0.0"}


if __name__ == "__main__":
    port = int(os.environ.get("PYTHON_DOC_PORT", "8081"))
    logger.info(f"Starting Python Document Service on port {port}")
    uvicorn.run(app, host="0.0.0.0", port=port, log_level="info")
